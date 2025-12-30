from reportlab.lib.pagesizes import letter
from docx import Document
from docx.shared import Inches
import os

def create_markdown(text, image_path, output_path, graph_path=None):
    """
    Creates a Markdown file with the analysis results.
    """
    try:
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write(f"# RxShield Analysis Report\n\n")
            f.write(f"**Image:** {image_path}\n\n")
            
            if graph_path and os.path.exists(graph_path):
                 f.write(f"## Knowledge Graph\n\n")
                 f.write(f"![Knowledge Graph]({graph_path})\n\n")
            
            f.write(f"## Analysis Results\n\n")
            f.write(text)
        return True
    except Exception as e:
        print(f"Error creating Markdown: {e}")
        return False

def create_word(text, image_path, output_path, graph_path=None):
    """
    Creates a Word document with the analysis results.
    """
    try:
        doc = Document()
        doc.add_heading('RxShield Analysis Report', 0)

        if image_path:
            try:
                doc.add_picture(image_path, width=Inches(6))
            except Exception as e:
                doc.add_paragraph(f"[Error adding image: {e}]")
        
        if graph_path and os.path.exists(graph_path):
            try:
                doc.add_heading('Knowledge Graph', level=2)
                doc.add_picture(graph_path, width=Inches(6))
            except Exception as e:
                doc.add_paragraph(f"[Error adding graph: {e}]")

        doc.add_heading('Analysis Results', level=1)
        doc.add_paragraph(text)

        doc.save(output_path)
        return True
    except Exception as e:
        print(f"Error creating Word doc: {e}")
        return False, f"Failed to save Word doc: {e}"

def create_pdf(text, image_path, output_path, graph_path=None):
    """
    Creates a professionally formatted PDF report using ReportLab.
    Parses the text to identify headers, lists, and bold content.
    """
    try:
        from reportlab.lib import colors
        from reportlab.lib.pagesizes import letter
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Image as RLImage, Table, TableStyle, PageBreak
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.units import inch
        
        doc = SimpleDocTemplate(output_path, pagesize=letter,
                                rightMargin=40, leftMargin=40,
                                topMargin=40, bottomMargin=40)
        story = []
        styles = getSampleStyleSheet()
        
        # --- Styles Definition ---
        title_style = ParagraphStyle(
            'RxTitle',
            parent=styles['Heading1'],
            fontSize=22,
            textColor=colors.HexColor('#008080'), # Teal
            alignment=1, # Center
            spaceAfter=20,
            fontName='Helvetica-Bold'
        )
        
        section_header = ParagraphStyle(
            'RxHeader',
            parent=styles['Heading2'],
            fontSize=14,
            textColor=colors.white,
            backColor=colors.HexColor('#2c3e50'), # Dark Blue
            borderPadding=(5, 10, 5, 10), # T, R, B, L
            spaceBefore=15,
            spaceAfter=10,
            fontName='Helvetica-Bold',
            keepWithNext=True
        )
        
        sub_header = ParagraphStyle(
            'RxSubHeader',
            parent=styles['Heading3'],
            fontSize=12,
            textColor=colors.HexColor('#2c3e50'),
            spaceBefore=10,
            spaceAfter=5,
            fontName='Helvetica-Bold'
        )
        
        body_style = ParagraphStyle(
            'RxBody',
            parent=styles['Normal'],
            fontSize=10,
            leading=14,
            spaceAfter=6,
            textColor=colors.black
        )
        
        disclaimer_style = ParagraphStyle(
            'RxDisclaimer',
            parent=styles['BodyText'],
            fontSize=9,
            textColor=colors.red,
            alignment=0,
            spaceBefore=20,
            fontName='Helvetica-Oblique'
        )

        footer_style = ParagraphStyle(
            'RxFooter',
            parent=styles['BodyText'],
            fontSize=8,
            textColor=colors.grey,
            alignment=1, # Center
            spaceBefore=30
        )

        # --- Content Generation ---

        # 1. Title
        story.append(Paragraph("RxShield Analysis Report", title_style))
        story.append(Spacer(1, 10))
        
        # 2. Date
        import datetime
        date_str = datetime.datetime.now().strftime("%B %d, %Y")
        story.append(Paragraph(f"Generated on: {date_str}", ParagraphStyle('Date', parent=body_style, alignment=1, textColor=colors.grey)))
        story.append(Spacer(1, 20))

        # 3. Images Section (Side by Side if possible, or stacked)
        # Using a Table for layout
        imgs_row = []
        
        # Scan Image
        if image_path and os.path.exists(image_path):
            try:
                # Constrain dimensions
                img = RLImage(image_path, width=3.2*inch, height=2.4*inch, kind='proportional')
                imgs_row.append([Paragraph("<b>Original Prescription</b>", styles['Normal']), img])
            except: pass
            
        # KG Image
        if graph_path and os.path.exists(graph_path):
            try:
                # Constrain dimensions
                img_kg = RLImage(graph_path, width=3.2*inch, height=3.0*inch, kind='proportional')
                # If we have both, we need a separate row or handle list
                # Let's stack them visually with labels for cleaner PDF flow usually
                # But user wants "Well Enhanced". Let's do clear sections.
            except: pass

        # Implementation: Sequential Images with Headers
        if image_path and os.path.exists(image_path):
            story.append(Paragraph("Scanned Prescription", sub_header))
            try:
                img = RLImage(image_path, width=5.5*inch, height=3.5*inch, kind='proportional')
                story.append(img)
            except: pass
            story.append(Spacer(1, 15))

        if graph_path and os.path.exists(graph_path):
            story.append(Paragraph("Knowledge Graph Visualization", sub_header))
            try:
                img_kg = RLImage(graph_path, width=6*inch, height=5*inch, kind='proportional')
                story.append(img_kg)
                story.append(Paragraph("<i>Visual representation of patient context and interactions.</i>", ParagraphStyle('Caption', parent=body_style, fontSize=8, alignment=1, textColor=colors.grey)))
            except: pass
            story.append(PageBreak())

        # 4. Text Analysis Content
        import re
        
        # Pre-process text to remove "Structure" noise if any
        lines = text.split('\n')
        
        current_section = None
        
        for line in lines:
            line = line.strip()
            if not line: continue
            
            # Detect Sections (=== ... ===) or Markdown Headers (# ..)
            if (line.startswith('=') and len(line) > 5) or (line.startswith('# ')):
                # Cleanup
                title = line.replace('=', '').replace('#', '').strip()
                if "Accuracy Score" in title: continue # Skip internal metric
                
                story.append(Paragraph(title, section_header))
            
            elif line.startswith('## '):
                subtitle = line.replace('#', '').strip()
                story.append(Paragraph(subtitle, sub_header))
                
            elif line.startswith('- ') or line.startswith('* '):
                # Bullet list
                # Parse bold syntax **text** -> <b>text</b>
                content = line[2:].strip()
                content = re.sub(r'\*\*(.*?)\*\*', r'<b>\1</b>', content)
                story.append(Paragraph(f"• {content}", ParagraphStyle('Bullet', parent=body_style, leftIndent=15, bulletIndent=0)))
                
            else:
                # Normal Text
                # Parse bold syntax
                content = line
                content = re.sub(r'\*\*(.*?)\*\*', r'<b>\1</b>', content)
                
                # Check for "Key-Value" pairs like "Patient Name:"
                if ':' in content and len(content.split(':')[0]) < 25:
                     parts = content.split(':', 1)
                     content = f"<b>{parts[0]}:</b>{parts[1]}"
                     
                story.append(Paragraph(content, body_style))

        # 5. Footer / Disclaimer
        story.append(Spacer(1, 30))
        story.append(Paragraph("DISCLAIMER: This analysis is generated by AI (RxShield). It is intended for assistance only and should NOT replace professional medical advice. Always verify with a certified pharmacist or doctor.", disclaimer_style))
        
        story.append(Paragraph("Powered by RxShield v4", footer_style))

        # Build
        doc.build(story)
        return True
        
    except Exception as e:
        print(f"Error creating PDF: {e}")
        return False
